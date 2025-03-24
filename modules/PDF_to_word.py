import streamlit as st
from pdf2docx import Converter
import os
import tempfile
import zipfile
from io import BytesIO
from utils.common import return_to_main
import base64
import shutil
from utils.common import cleanup_temp_dirs

# 尝试导入python-docx库
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 尝试导入PyPDF2库用于获取PDF页数
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

def convert_pdf_to_docx(pdf_path, docx_path):
    """将单个PDF文件转换为Word文档"""
    try:
        # 首先尝试常规转换方式
        cv = Converter(pdf_path)
        cv.convert(
            docx_path,
            start=0,
            end=None,
            pages=None,
            zoom=1.5,
            multi_processing=True,
            grayscale=False,
            use_cropbox=True,
            ignore_errors=True
        )
        cv.close()
        return True, None  # 成功标志和跳过页面列表
    except Exception as e:
        # PNG颜色空间问题时，尝试更保守的设置
        if "unsupported colorspace for 'png'" in str(e):
            try:
                cv = Converter(pdf_path)
                cv.convert(
                    docx_path,
                    start=0,
                    end=None,
                    pages=None,
                    zoom=1,
                    multi_processing=False,
                    grayscale=True,
                    use_cropbox=True,
                    ignore_errors=True
                )
                cv.close()
                return True, None
            except Exception as e2:
                # 如果仍然失败，尝试逐页转换，跳过问题页面
                # 检查是否有python-docx库
                if not DOCX_AVAILABLE:
                    raise Exception("处理特殊格式需要python-docx库支持。请联系管理员安装该库。")
                    
                try:
                    # 获取总页数
                    total_pages = get_pdf_page_count(pdf_path)
                    
                    # 创建临时文件存储各页内容
                    temp_docs = []
                    skipped_pages = []
                    
                    # 创建页面转换的子进度条
                    page_progress = st.progress(0)
                    page_status = st.empty()
                    page_status.text(f"正在逐页转换 '{os.path.basename(pdf_path)}': 第0/{total_pages}页")
                    
                    # 逐页转换
                    for page_num in range(total_pages):
                        # 更新页面转换进度
                        page_progress.progress((page_num) / total_pages)
                        page_status.text(f"正在逐页转换 '{os.path.basename(pdf_path)}': 第{page_num+1}/{total_pages}页")
                        
                        temp_docx = f"{pdf_path}_{page_num}.docx"
                        try:
                            cv = Converter(pdf_path)
                            # 尝试获取当前页，确保该页存在
                            try:
                                # 安全检查：确认该页面真实存在
                                if cv.store.get_page_count() <= page_num:
                                    print(f"页面 {page_num + 1} 超出文档范围，跳过")
                                    skipped_pages.append(page_num + 1)
                                    cv.close()
                                    continue
                            except:
                                pass  # 如果检查失败，仍然尝试转换
                                
                            cv.convert(
                                temp_docx,
                                start=page_num,
                                end=page_num+1,
                                pages=[page_num],
                                zoom=1.0,  # 调整缩放比例为标准大小
                                multi_processing=False,
                                grayscale=True,
                                line_spacing=1.0,  # 设置行间距
                                use_cropbox=True,  # 使用裁剪框而不是媒体框
                                adjust_tables=True,  # 优化表格处理
                                ignore_errors=True
                            )
                            cv.close()
                            temp_docs.append(temp_docx)
                        except Exception as e3:
                            # 记录跳过的页面
                            skipped_pages.append(page_num + 1)  # 转为显示给用户的页码(从1开始)
                            # 记录错误但继续处理
                            print(f"跳过页面 {page_num + 1}: {str(e3)}")
                            continue
                    
                    # 完成页面转换进度
                    page_progress.progress(1.0)
                    page_status.text(f"正在合并转换的页面 '{os.path.basename(pdf_path)}'")
                    
                    # 如果没有任何页面成功转换
                    if not temp_docs:
                        raise Exception("所有页面转换均失败")
                        
                    # 合并成功转换的页面
                    from docx.shared import Pt, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    # 创建一个新的空白文档
                    merged_doc = docx.Document()
                    
                    # 设置文档页面属性，避免不必要的大空白边距
                    for section in merged_doc.sections:
                        section.page_width = Inches(8.5)
                        section.page_height = Inches(11)
                        section.left_margin = Inches(1)
                        section.right_margin = Inches(1)
                        section.top_margin = Inches(1)
                        section.bottom_margin = Inches(1)
                    
                    successful_merges = 0
                    
                    for i, temp_doc_path in enumerate(temp_docs):
                        try:
                            temp_doc = docx.Document(temp_doc_path)
                            
                            # 只对第二页及之后的页面添加分页符
                            if i > 0:
                                run = merged_doc.add_paragraph().add_run()
                                run.add_break(docx.enum.text.WD_BREAK.PAGE)
                            
                            # 复制所有内容元素，但跳过空段落和不必要的分隔符
                            for element in temp_doc.element.body:
                                # 忽略完全空白的段落和不必要的分节符
                                if element.tag.endswith('p') and not element.text_content().strip():
                                    continue
                                
                                # 忽略可能包含过多空白的某些元素
                                if element.tag.endswith('sectPr'):
                                    continue
                                    
                                try:
                                    # 复制文档元素
                                    merged_doc.element.body.append(element)
                                    successful_merges += 1
                                except Exception as element_error:
                                    print(f"复制元素失败: {str(element_error)}")
                                    continue  # 继续处理下一个元素
                        except Exception as merge_error:
                            print(f"合并文档出错 {temp_doc_path}: {str(merge_error)}")
                        finally:
                            # 删除临时文件
                            try:
                                os.unlink(temp_doc_path)
                            except:
                                pass
                    
                    # 合并成功转换的页面后清理文档
                    def clean_document(doc):
                        """删除多余的空白段落和修复格式问题"""
                        # 找出并删除连续的空白段落（仅保留一个）
                        i = 0
                        while i < len(doc.paragraphs) - 1:
                            if not doc.paragraphs[i].text.strip() and not doc.paragraphs[i+1].text.strip():
                                p = doc.paragraphs[i]._element
                                p.getparent().remove(p)
                                # 删除后，不增加索引，因为集合大小已经减少了
                            else:
                                i += 1
                                
                        # 修复段落格式（设置适当的字体和行距）
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():  # 只处理非空段落
                                for run in paragraph.runs:
                                    # 设置默认字体和大小
                                    if not run.font.size:
                                        run.font.size = Pt(11)
                                    if not run.font.name:
                                        run.font.name = 'Arial'
                    
                    # 执行文档清理
                    try:
                        clean_document(merged_doc)
                    except Exception as clean_error:
                        print(f"清理文档出错: {str(clean_error)}")
                    
                    # 保存合并后的文档
                    merged_doc.save(docx_path)
                    
                    # 清理页面转换的进度显示
                    page_progress.empty()
                    page_status.empty()
                    
                    if skipped_pages:
                        return False, skipped_pages  # 部分成功，返回跳过的页面列表
                    else:
                        return True, None  # 全部页面转换成功
                        
                except Exception as e4:
                    # 如果页面处理方法也失败，抛出原始异常
                    raise Exception(f"转换失败，PDF包含不支持的PNG图像格式。尝试逐页转换也失败。错误信息：{str(e)}")
        
        # 非PNG颜色空间问题的其他错误
        raise Exception(f"转换失败，请检查PDF文件格式。错误信息：{str(e)}")


def create_zip_file(file_paths):
    """创建包含所有转换后文件的zip文件"""
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file_path, original_name in file_paths:
            with open(file_path, 'rb') as f:
                zip_file.writestr(original_name.replace('.pdf', '.docx'), f.read())
    return zip_buffer.getvalue()


def pdf_to_word():
    return_to_main()
    st.title("PDF➡️Word")
    st.write("上传多个PDF文件，将其转换为Word文档")
    
    # 初始化临时目录跟踪
    if 'temp_dirs' not in st.session_state:
        st.session_state.temp_dirs = []
    
    # 显示可能的兼容性提示
    with st.expander("PDF转换兼容性提示", expanded=False):
        st.info("""
        **转换提示：**
        1. 如果PDF包含特殊格式的PNG图像，可能会出现"unsupported colorspace for 'png'"错误
        2. 系统会尝试多种方法处理这类文件，但可能会导致部分页面被跳过
        3. 如果遇到转换问题，可以尝试先用其他工具（如Adobe Acrobat）重新保存PDF后再转换
        """)
    
    # 允许多文件上传
    uploaded_files = st.file_uploader("选择PDF文件（可多选）", 
                                    type=['pdf'], 
                                    accept_multiple_files=True)
    
    if uploaded_files:
        st.write(f"已上传 {len(uploaded_files)} 个文件")
        
        if st.button('开始转换'):
            # 清理之前的临时目录
            cleanup_temp_dirs()
            
            # 创建新的临时目录
            temp_dir = tempfile.mkdtemp()
            st.session_state.temp_dirs.append(temp_dir)
            
            converted_files = []
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                for index, uploaded_file in enumerate(uploaded_files):
                    # 更新进度
                    progress = (index) / len(uploaded_files)
                    progress_bar.progress(progress)
                    status_text.text(f"正在转换: {uploaded_file.name}")
                    
                    # 创建临时PDF文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                        tmp_pdf.write(uploaded_file.getvalue())
                        pdf_path = tmp_pdf.name
                    
                    # 创建临时Word文件路径
                    docx_path = pdf_path.replace('.pdf', '.docx')
                    
                    try:
                        # 执行转换
                        success, skipped_pages = convert_pdf_to_docx(pdf_path, docx_path)
                        if success:
                            converted_files.append((docx_path, uploaded_file.name))
                            st.success(f"'{uploaded_file.name}' 转换成功！")
                        else:
                            # 部分成功的情况
                            converted_files.append((docx_path, uploaded_file.name))
                            # 以通知形式提示用户某些页面被跳过
                            if len(skipped_pages) > 0:
                                if len(skipped_pages) == 1:
                                    page_str = str(skipped_pages[0])
                                    st.warning(f"'{uploaded_file.name}' 转换时第 {page_str} 页出现问题被跳过。该页可能包含不支持的PNG图像格式。")
                                else:
                                    # 如果跳过的页面连续，则显示范围
                                    if len(skipped_pages) > 3 and skipped_pages[-1] - skipped_pages[0] + 1 == len(skipped_pages):
                                        page_str = f"{skipped_pages[0]}-{skipped_pages[-1]}"
                                    else:
                                        page_str = ", ".join(map(str, skipped_pages))
                                    st.warning(f"'{uploaded_file.name}' 转换时第 {page_str} 页出现问题被跳过。这些页面可能包含不支持的PNG图像格式。")
                                
                                # 显示如何处理这种情况的提示
                                if index == len(uploaded_files) - 1:  # 只在最后一个文件时显示
                                    st.info("提示：您可以用其他工具(如Adobe Acrobat)打开PDF，重新保存为PDF后再尝试转换，可能会解决部分页面无法转换的问题。")
                    except Exception as e:
                        # 过滤掉特定错误消息，避免显示技术细节
                        error_msg = str(e)
                        if "unsupported colorspace for 'png'" in error_msg:
                            st.error(f"转换 '{uploaded_file.name}' 失败: 文件包含不支持的PNG图像格式")
                        else:
                            st.error(f"转换 '{uploaded_file.name}' 失败，已跳过该文件")
                    finally:
                        # 删除临时PDF文件
                        os.unlink(pdf_path)
                
                # 完成进度条
                progress_bar.progress(1.0)
                status_text.text("转换完成！")
                
                if converted_files:
                    if len(converted_files) == 1:
                        # 单个文件直接提供下载
                        docx_path, original_name = converted_files[0]
                        with open(docx_path, 'rb') as f:
                            docx_data = f.read()
                        
                        st.success('文件转换成功！')
                        
                        # 添加CSS样式
                        st.markdown("""
                        <style>
                        .download-button {
                            display: inline-block;
                            padding: 8px 16px;
                            background-color: #4CAF50;
                            color: white !important;
                            text-align: center;
                            text-decoration: none;
                            font-size: 16px;
                            margin: 10px 5px;
                            border-radius: 4px;
                            cursor: pointer;
                            transition: background-color 0.3s;
                        }
                        .download-button:hover {
                            background-color: #45a049;
                        }
                        </style>
                        """, unsafe_allow_html=True)
                        
                        # 创建base64编码的Word文档下载链接
                        docx_filename = original_name.replace('.pdf', '.docx')
                        b64_docx = base64.b64encode(docx_data).decode()
                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        href = f'<a href="data:{mime_type};base64,{b64_docx}" download="{docx_filename}" class="download-button">下载Word文档</a>'
                        st.markdown(href, unsafe_allow_html=True)
                    else:
                        # 多个文件创建zip文件
                        zip_data = create_zip_file(converted_files)
                        
                        # 提供zip文件下载
                        if len(converted_files) < len(uploaded_files):
                            st.success(f'成功转换 {len(converted_files)}/{len(uploaded_files)} 个文件！')
                            st.info('部分文件转换失败，已跳过。已成功的文件可以下载。')
                        else:
                            st.success(f'成功转换 {len(converted_files)} 个文件！')
                        
                        # 添加CSS样式
                        st.markdown("""
                        <style>
                        .download-button {
                            display: inline-block;
                            padding: 8px 16px;
                            background-color: #4CAF50;
                            color: white !important;
                            text-align: center;
                            text-decoration: none;
                            font-size: 16px;
                            margin: 10px 5px;
                            border-radius: 4px;
                            cursor: pointer;
                            transition: background-color 0.3s;
                        }
                        .download-button:hover {
                            background-color: #45a049;
                        }
                        </style>
                        """, unsafe_allow_html=True)
                        
                        # 创建base64编码的ZIP下载链接
                        b64_zip = base64.b64encode(zip_data).decode()
                        href = f'<a href="data:application/zip;base64,{b64_zip}" download="converted_documents.zip" class="download-button">下载所有Word文档</a>'
                        st.markdown(href, unsafe_allow_html=True)
                else:
                    st.error("所有文件转换均失败。请检查PDF文件格式或尝试其他工具处理。")
                
            except Exception as e:
                error_msg = str(e)
                if "unsupported colorspace for 'png'" in error_msg:
                    st.error("处理过程中出现PNG图像格式错误，请尝试用其他工具编辑PDF后再试。")
                else:
                    st.error("处理过程中出现错误，请重试或联系管理员。")
            
            finally:
                # 清理所有临时文件
                for docx_path, _ in converted_files:
                    try:
                        os.unlink(docx_path)
                    except:
                        pass

# 获取PDF总页数的函数
def get_pdf_page_count(pdf_path):
    """使用多种方法获取PDF总页数"""
    # 优先使用PyPDF2获取页数
    if PYPDF2_AVAILABLE:
        try:
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                pages = len(pdf_reader.pages)
                if pages > 0:
                    return pages
        except Exception as e:
            print(f"PyPDF2获取页数失败: {str(e)}")
            pass  # 如果失败，继续尝试pdf2docx方法
    
    # 使用pdf2docx方法获取页数
    try:
        temp_cv = Converter(pdf_path)
        total_pages = temp_cv.store.get_page_count()
        temp_cv.close()
        if total_pages > 0:
            return total_pages
    except Exception as e:
        print(f"Converter获取页数失败: {str(e)}")
    
    # 如果都失败了，返回一个默认值
    return 1  # 至少尝试转换一页
